# -*- coding: utf-8 -*-
"""Generates F:\AzureCore\AzureDoc\AzureBranches-26.1.2-B_LINEAR_V4-INTRO.docx
b_linear_v4 storage backend standalone introduction (docx counterpart of
B_LINEAR_V4-INTRO.md). Style mirrors the EXP document series."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"F:\AzureCore\AzureDoc\AzureBranches-26.1.2-B_LINEAR_V4-INTRO.docx"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

BOLD_RE = re.compile(r"(\*\*.*?\*\*)")
CODE_RE = re.compile(r"(`[^`]*`)")


def _set_run(run, size, bold, font_east="宋体", ascii_font="Times New Roman"):
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_east)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)


def para(text, size=10.5, bold_all=False, align="justify", indent_chars=0, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    if indent_chars > 0:
        p.paragraph_format.first_line_indent = Pt(size * indent_chars)
    for part in BOLD_RE.split(text):
        if not part:
            continue
        m = BOLD_RE.fullmatch(part)
        if m:
            _set_run(p.add_run(m.group(1)[2:-2]), size, True)
            continue
        for sub in CODE_RE.split(part):
            if not sub:
                continue
            if CODE_RE.fullmatch(sub):
                _set_run(p.add_run(sub[1:-1]), size, False, ascii_font="Consolas")
            else:
                _set_run(p.add_run(sub), size, bold_all)
    return p


def code(text, size=8.5, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    _set_run(r, size, False, ascii_font="Consolas")
    return p


def table(rows, widths=None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = t.cell(i, j)
            p = cell.paragraphs[0]
            r = p.add_run(str(cell_text))
            _set_run(r, 9.5, i == 0)
            if i == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "E7E6E6")
                cell._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()


# ---------------- Title block ----------------
para("AzureBranches", 22, True, "center", 0, 0, 6)
para("b_linear_v4 介绍文档", 16, True, "center", 0, 6, 0)
para("区域文件存储后端：从「b_linear 移植」到「VERSION 4 + 校验链 + 配置开关（默认 MCA）」的完整介绍", 12, False, "center", 0, 6, 0)
para("版本 26.1.2-EXP7　　2026 年 08 月 22 日", 10.5, False, "center", 0, 6, 0)
para("基于 Folia Regionized Ticking 模型", 10.5, False, "center", 0, 6, 0)
para("感谢 PaperMC / Folia (Spottedleaf) / LuminolMC (EarthMe 及其团队) / Arbor (Little 及 xymb 血统)", 10.5, False, "center", 0, 12, 0)

para("配套文档：**STORAGE-V4-SPEC.md**（v4 格式正式规范：字节级布局/常量/兼容矩阵/校验定义）；**AzureBranches-26.1.2-EXP7.docx**（EXP7 工程报告：集成方法/缺陷修复记录/验证数据）。本文是概念、架构与使用介绍，与两者互补，不重复细节。", 9.5, False, "justify", 0, 8, 0)

# ---------------- 1 ----------------
para("一、这是什么", 14, True, "left", 0, 6, 12)
para("**b_linear_v4 是 AzureBranches 提供的一种可选的区域文件存储后端**，用于区域文件（region/r.0.0.mca、实体、POI 等）：(1) 移植自 Luminol/Arbor 的 b_linear 存储体系（MrHua269 / Little / xymb 血统，GPLv3 派生，见 NOTICE.md）；(2) 写侧使用本项目的 **VERSION 4** 主文件布局，读侧兼容上游 v0x02（整文件 ZSTD）/v0x03（桶式）以及 xymb 祖宗线性格式；(3) 经配置开关 storage.region_format 选择，**默认仍是 vanilla MCA**——不配置就是原样行为，零依赖零漂移。", 10.5, False, "justify", 2, 6, 0)
para("一句话：**b_linear_v4 = 更少的写放大 + 更快的写入 + 带完整校验链的文件格式 + 崩溃后 master 永不被污染。**", 10.5, False, "justify", 2, 6, 0)

# ---------------- 2 ----------------
para("二、为什么需要它（对照 vanilla MCA）", 14, True, "left", 0, 6, 12)
table([
    ["维度", "vanilla MCA", "b_linear_v4"],
    ["每 chunk 写入", "全量重写 8 KiB 定位表（1024 槽 × 12B），chunk 数据在 4 KiB 扇区内就地覆盖", "数据追加进 swap 工作文件；master 只在回写窗口（写后空闲 3s / 关闭 / 首写）整体快照一次"],
    ["文件整体校验", "仅头部一个摘要，chunk 数据无独立校验", "四层校验链：位置表 xxhash64（footer）→ 桶块 xxhash32 → 段长上界 → chunkSection 内部 xxhash32"],
    ["损坏感知", "损坏可能静默读出坏数据", "任一层失败 → IOException 带文件路径与偏移，绝不返回部分数据"],
    ["写失败", "存在静默路径", "同步失败在 flusher 日志显式报告；加载失败后拒绝把空工作态覆写回 master（防空写破坏）"],
    ["大文件", "Oversized 需要附属文件（.oversized）", "桶内带 64 位时间戳的 section 结构，超限在写入点显式 RegionFileSizeException"],
    ["启动加载", "全文件按需扇区读", "按桶惰性加载（16 桶 × 64 chunk），只读用到的桶"],
])

# ---------------- 3 ----------------
para("三、术语表", 14, True, "left", 0, 6, 12)
table([
    ["术语", "含义"],
    ["master", "主文件 r.<rx>.<rz>.mca——已提交快照，v4 布局，仅由同步窗口写入"],
    ["swap", "交换文件 r.<rx>.<rz>.mca.swp——写工作文件，LZ4 段、1024 个 Sector、DELETE_ON_CLOSE"],
    ["Sector", "swap 内的 1024 个槽位，每个对应一个 chunk 位置（17B 表项：offset/length/hasData + 数据区）"],
    ["Bucket", "逻辑桶 16 个 × 每桶 64 个 chunk 槽（BUCKET_SHIFT=6），master 压缩/存储/校验的粒度"],
    ["写纪元", "每桶 writeEpoch/syncedEpoch 两个原子计数：脏桶 = 写纪元 ≠ 已同步纪元"],
    ["chunkSection", "单 chunk 的存储单元：len(4) | timestamp(8) | xxhash32(4) | data"],
    ["回写窗口", "距最后一次写入超过 flushOfWriteTimeoutMs（默认 3000ms）即触发一次 master 同步"],
])

# ---------------- 4 ----------------
para("四、架构总览", 14, True, "left", 0, 6, 12)
code("""             ┌───────────────────────────── 一次运行 ─────────────────────────────┐
 打开 region │ 删除旧 .swp → 建空 .swp → 惰性加载 master 桶（用到的桶才读）          │
             ▼                                                                     │
   chunk 写  ├─→ swap：LZ4 压缩 → 追加到 Sector 数据区（不碰 master）               │
             │     桶写纪元++ → 记录 LAST_WRITTEN                                  │
   chunk 读  ├─→ 直接读 swap Sector；swap 没有的桶才从 master 惰性加载进 swap         │
             ▼                                                                     │
  回写窗口   ├─→ 脏桶序列化：rawLen|compressedLen|ZSTD → master.tmp                 │
  或 close   │     位置表 + footer（xxhash64）→ master.tmp                          │
             │     tmp → ATOMIC_MOVE → master（替换失败自动退回非原子 move）         │
             ▼                                                                     │
 进程结束    └─→ swap 随句柄关闭被删除；master 出厂。                               │
             ▲                                                                     │
  硬杀(崩溃) └─→ 同上：swap 无残留（DELETE_ON_CLOSE），master 保持最后一次同步状态   │""")
para("swap 是工作台，master 是成品：所有写操作只碰 swap；master 只在回写窗口被整体重建，且永远先写 .tmp 再原子替换。", 10.5, False, "justify", 2, 2, 0)
para("一个 flusher 管全部文件：进程级单例（RegionFormat.flusher()），4 个 IO 线程 + 20ms 检查器 + 3s 写后超时，负责每个文件的 syncIfNeeded 调度与背压。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 5 ----------------
para("五、写入路径：chunk 如何落盘", 14, True, "left", 0, 6, 12)
para("1. Moonrise 写到 getChunkDataOutputStream(pos)（即 ChunkBufferHelper）：组装 chunkSection = len(4) | timestamp(8) | xxhash32(4) | data；writeChunkDataRaw 做 LZ4 压缩 → sector.store() 追加到 swap → 桶写纪元递增；关闭流时触发 flushInternal()（自动压缩检查：swap 空闲 > 1 MiB 且 > 60% 时压缩整理）。", 10.5, False, "justify", 2, 3, 0)
para("2. **回写窗口**：距上次写超过 3s（flushOfWriteTimeoutMs），flusher 对该文件执行 syncToMasterFile：脏桶（写纪元 ≠ 已同步纪元）从 swap 读数据 → ZSTD 压缩 → 写入 master.tmp；**未脏桶整块搬运**——直接从旧 master（v3/v4）按字节复制（v3→v4 迁移免重压缩）；写位置表（每项：offset | compressedLen | 桶块 xxhash32）与 footer（位置表 xxhash64）；master.tmp → 原子替换 master，同时标记各桶已同步纪元。", 10.5, False, "justify", 2, 3, 0)
para("3. **关闭**（优雅停服 / 区域文件关页）：syncIfNeeded + 关闭 swap，数据全部进 master。", 10.5, False, "justify", 2, 3, 0)
para("换文件后已提交数据仍是最后一次同步的数据；同步周期内新写入在 swap 里，进程正常关闭时全部落盘。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 6 ----------------
para("六、master v4 布局（速览）", 14, True, "left", 0, 6, 12)
code("""[0,     14)  header     superblock(long) | version(0x04) | compressionLevel(byte) | xxHash32Seed(int)
[14,   270)  位置表     16 × 16B：bucketOffset(long) | compressedLen(int) | bucketXXHash32(int)
                       （全 0 = 该桶无数据；offset 必须落在数据区）
[270,  EOF)  桶数据区   脏桶按序：rawLen(int) | compressedLen(int) | ZSTD(compressedLen)
                       ZSTD 解压后为 64 个槽位：secLen(int) | chunkSection（§5-1 结构）
[EOF-21,EOF) footer     superblock(long) | version(0x04) | positionTableXXHash64(long) | reserved(4)""")
para("要点：", 10.5, False, "left", 2, 2, 0)
para("· 桶哈希覆盖**整个桶块**（rawLen + compressedLen + ZSTD 载荷），不解压即验；", 10.5, False, "justify", 2, 2, 0)
para("· footer 表哈希覆盖整张 256 B 位置表；", 10.5, False, "justify", 2, 2, 0)
para("· 读取校验链：超块 → 版本分派 → 位置表（边界 + footer 哈希）→ 桶（长度上界 + 桶哈希）→ 段（长度上界）→ chunkSection（内部 xxhash32）。任何一层失败 → IOException，绝不返回部分数据。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 7 ----------------
para("七、崩溃与恢复语义", 14, True, "left", 0, 6, 12)
para("· **swap 是 DELETE_ON_CLOSE**：进程无论正常退出还是硬杀，swap 都会被系统删除，不会残留半成品 WAL。", 10.5, False, "justify", 2, 2, 0)
para("· **崩溃恢复点 = 最后一次 master 同步**：硬杀后 master 保持上一次同步的内容，永不被部分写或空写污染（master 永远原子替换 + 分段可验）。", 10.5, False, "justify", 2, 2, 0)
para("· **防空写**：若本次运行加载 master 时发生任何校验失败（loadFailed），后续 syncToMasterFile 拒绝执行——防止 try-with-resources 关闭路径把空工作态写坏好的 master（EXP7 实测捕获过该级联，已修复）。", 10.5, False, "justify", 2, 2, 0)
para("· 含义：写后 3s 内未同步的写可能因崩溃丢失（与 vanilla 缓冲写窗口同量级）；但**绝不会损坏已提交数据**。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 8 ----------------
para("八、配置与启用", 14, True, "left", 0, 6, 12)
code("""[storage]
region_format = "b_linear_v4"      # "mca"（默认）| "b_linear_v4"

[storage.linear]
compression_level = 1              # 1..22，ZSTD 压缩级别（默认 1）""")
para("· 启动日志确认行：[AzureBranches] storage.region_format=b_linear_v4 (compression 1)；", 10.5, False, "justify", 2, 2, 0)
para("· 配置在**首次区域文件创建前**应用（RegionFormat.open 内 RegionFormatBootstrap.ensureApplied()，幂等）；配置未加载时自动重试，非法值一次性告警并回退 MCA；", 10.5, False, "justify", 2, 2, 0)
para("· **切换须知**：b_linear_v4 不能读 vanilla MCA 区域文件（见第九节），老世界启用前请先备份或按 xymb 迁移路径处理；建议在新世界上启用。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 9 ----------------
para("九、兼容矩阵 / 迁移", 14, True, "left", 0, 6, 12)
table([
    ["主文件", "版本", "处理"],
    ["superblock=-0x200812250269L", "0x02", "旧整文件 ZSTD → 迁移入 swap → 下次同步升 v4（读路径保留，尚未生成样本实测）"],
    ["同上", "0x03", "旧桶式 → 按桶惰性加载迁移 → 升 v4（已用真实上游引擎样本实测）"],
    ["同上", "0x04", "本格式"],
    ["superblock=0xc3ff13183cca9d9aL", "1/2/3", "xymb 祖宗线性 → 逐块迁移（保留）"],
    ["其他", "—", "硬错误拒绝（带超块十六进制与路径）"],
])
para("· **v3 → v4 就地迁移**：打开旧文件 → 正常读写 → 下次同步自动写出 v4（未改动桶整块字节搬运，不重压缩）。", 10.5, False, "justify", 2, 2, 0)
para("· **MCA ↔ b_linear_v4 不互通**：这是设计边界（v4 只兼容 xymb 系）。切格式 = 换世界或先行迁移。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 10 ----------------
para("十、性能（EXP7 实测基线）", 14, True, "left", 0, 6, 12)
para("JVM 级微基准（同进程顺序写，64 × 64 KiB chunk，含 flush）：", 10.5, False, "justify", 2, 2, 0)
table([
    ["后端", "写 + flush", "主文件体积"],
    ["b_linear_v4", "44.8 MiB/s", "4.20 MB"],
    ["vanilla MCA", "12.5 MiB/s", "4.46 MB"],
])
para("≈ 3.6× 写入吞吐，体积略优。真实服务器场景（多区域并发、长时间混合读写）尚未建立压力基线。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 11 ----------------
para("十一、源码地图", 14, True, "left", 0, 6, 12)
para("覆盖范围：chunk（region/）、实体（entities/）、POI（poi/）三类区域文件都经 RegionFormat.open 创建——实体/POI 存储同样是 RegionFileStorage 子类，自动获得同样的 v4 与校验语义。", 10.5, False, "justify", 2, 2, 0)
table([
    ["文件", "职责"],
    ["folia-server/azurepatches-new/com/azurebranches/storage/BufferedLinearRegionFile.java", "主引擎：swap/master/桶/纪元/校验链/v4 解析与写出"],
    [".../BufferedLinearRegionFileFlusher.java", "进程级保存调度器（IO 池 + 检查器 + 背压）"],
    [".../IRegionFile.java", "后端抽象（extends ChunkSystemRegionFile）"],
    [".../RegionFormat.java", "格式选择器（MCA / B_LINEAR_V4）+ 静态 open 创建点"],
    [".../RegionFormatBootstrap.java", "配置应用引导（在 open 处惰性执行）"],
    [".../net/minecraft/world/level/chunk/storage/RegionFileAdapter.java", "vanilla MCA 适配器（公共化受保护方法）"],
    ["folia-server/azurepatches-src/net/minecraft/world/level/chunk/storage/RegionFileStorage.java", "覆盖层：RegionFile → IRegionFile，创建点收敛到 RegionFormat.open"],
    ["folia-server/build.gradle.kts", "transformSource 锚点 + zstd-jni / zero-allocation-hashing 依赖（lz4 已在运行时类路径）"],
    ["azurebranches-common/.../AzureBranchesConfig.java", "storage.region_format / storage.linear.compression_level 配置项"],
    ["STORAGE-V4-SPEC.md", "v4 正式字节级规范"],
])

# ---------------- 12 ----------------
para("十二、已知限制与后续", 14, True, "left", 0, 6, 12)
para("1. 崩溃恢复点 = 最后一次 master 同步（swap 不保留）；flushOfWriteTimeoutMs（默认 3s）窗口内的未同步写可能丢失。", 10.5, False, "justify", 2, 3, 0)
para("2. MCA ↔ b_linear_v4 不互通；老世界启用需备份/迁移，建议新世界启用。", 10.5, False, "justify", 2, 3, 0)
para("3. v0x02 读路径保留但未生成样本实测（上游当前版本写 v0x03）。", 10.5, False, "justify", 2, 3, 0)
para("4. 性能为单机微基准，未覆盖真负载压力曲线。", 10.5, False, "justify", 2, 3, 0)
para("后续方向（见 EXP7 报告 §6.2）：v0x02 样本实测、MCA→v4 单区域迁移工具、服务端压力基线、同步重试与统计出口、可选持久 WAL（与上游语义偏离，需评估）。", 10.5, False, "justify", 2, 6, 0)

para("文档：AzureBranches-26.1.2-B_LINEAR_V4-INTRO　　版本：26.1.2-EXP7　　日期：2026 年 08 月 22 日　　项目：AzureBranches (https://github.com/XCxyTianQ/AzureBranches)", 9, False, "center", 0, 12, 0)

doc.save(OUT)
print("DOCX written:", OUT)
