# -*- coding: utf-8 -*-
"""Generates F:\AzureCore\AzureDoc\AzureBranches-26.1.2-EXP6.docx
Format mirrors the EXP5Plus document (direct formatting, CN section numbering)."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"F:\AzureCore\AzureDoc\AzureBranches-26.1.2-EXP6.docx"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

BOLD_RE = re.compile(r"(\*\*.*?\*\*)")


def _set_run(run, size, bold, font_east="宋体"):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_east)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def para(text, size=10.5, bold_all=False, align="justify", indent_chars=0, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    if indent_chars > 0:
        p.paragraph_format.first_line_indent = Pt(size * indent_chars)
    for part in BOLD_RE.split(text):
        if not part:
            continue
        m = BOLD_RE.fullmatch(part)
        if m:
            _set_run(p.add_run(m.group(1)[2:-2]), size, True)
        else:
            _set_run(p.add_run(part), size, bold_all)
    return p


def table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = t.cell(i, j)
            p = cell.paragraphs[0]
            r = p.add_run(str(cell_text))
            _set_run(r, 9.5, i == 0)
            if i == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "E7E6E6")
                cell._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()


# ---------------- Title block ----------------
para("AzureBranches", 22, True, "center", 0, 0, 6)
para("EXP6：EntityLayer 注入与 /data 实体 NBT 的 OCC 闭环", 16, True, "center", 0, 6, 0)
para("从“命令层捕获点选择”到“读集校验 + 按实体区域补偿”的实体数据一致性工程", 12, False, "center", 0, 6, 0)
para("版本 26.1.2-EXP6　　2026 年 08 月 16 日", 10.5, False, "center", 0, 6, 0)
para("基于 Folia Regionized Ticking 模型", 10.5, False, "center", 0, 6, 0)
para("感谢 PaperMC / Folia (Spottedleaf) / LuminolMC (EarthMe 及其团队)", 10.5, False, "center", 0, 12, 0)

# ---------------- Abstract ----------------
para("**摘要**", 10.5, False, "justify", 2, 6, 0)
para("EXP6（v26.1.2-EXP6）完成 EntityLayer 的注入，使实体 NBT 成为继方块、积分板之后第三个完整走通“捕获 → 验证 → 补偿”的 OCC 数据域。捕获点选择于**命令层**而非数据池：DataCommands 覆盖层在 /data 的六条执行流（get/getNumeric/resolveSourcePath/modify/merge/remove）中以 NBT 路径粒度记录读集取值（**recordReadValue**）与写前旧值（**interceptWrite**，含 REMOVED 删除哨兵），EntityLayer 新增 **parsePathString** 将 26.1 的路径文法（向量下标、equipment 槽位、复合子字段）归一到稳定复合键；PhaseSnapshot 的四个 NBT 映射改为并发容器以容忍实体区域线程上的捕获。校验侧：verifyReadSetAndResume 按实体分组、经 **onEntityAsync** 在实体所属区域重读实况并与观测值比对，接通 **CHECK_NBT_READ_SET** 五参 validate；补偿侧修正 EXP5Plus 的 P2 潜在缺陷——EntityLayer 补偿从全局 tick 线程改道**按实体分派的 onEntityAsync**（**compensateFor**）。工程中发现并修复 Folia 实体调度器的一个隐蔽语义陷阱：scheduleOrExecute 在实体所属 tick 线程上仍会排队到下一 tick，导致 EXP5 的“同区域快路径”在命令方块链内从未生效——RegionCommandExecutor.onEntityAsync 改为在 TickThread.isTickThreadFor(entity) 时内联执行，链内 /data 读写恢复同步快路径。全文以“线程模型 — 捕获点选择 — 键归一 — 控制流 — 一致性 — 验证”为主线，并以命令方块链内读己写、不可重复读、Δ/恢复补偿与重放收敛的实测数据佐证。", 10.5, False, "justify", 2, 6, 0)
para("**关键词**：Folia；命令方块；EntityLayer；/data；NBT 路径；PhaseSnapshot；OCC；CHECK_NBT_READ_SET；compensateFor；scheduleOrExecute；isTickThreadFor；26.1 NBT 键名迁移", 10.5, False, "justify", 0, 12, 0)

# ---------------- 一 ----------------
para("一、背景与问题定义", 14, True, "left", 0, 6, 12)
para("1.1　EXP5Plus 遗留：EntityLayer 捕获端零调用", 12, True, "left", 0, 6, 6)
para("EXP4 声明的实体 NBT 层在 EXP5Plus 结束时仍只有补偿端骨架：rollbackAndRetryExpChain 经构建期注入调用 EntityLayer.compensate，但 **interceptRead / interceptWrite 在全部生成源码中零调用点**——nbtOldValues 恒为空，补偿恒为 no-op；verifyReadSetAndResume 的 validate 调用中 NBT 读集参数缺失，CHECK_NBT_READ_SET 分支不可达。EXP5Plus 文档 7.2 节将“EntityLayer 捕获端接线”列为 P0，即本版本的工作起点。", 10.5, False, "justify", 2, 6, 0)
para("1.2　26.1 的 NBT 键名迁移：捕获与补偿的字典必须换新", 12, True, "left", 0, 6, 6)
para("Minecraft 26.1（1.21.6 系列）对实体 NBT 进行**部分 snake_case 迁移**，且装备栏发生结构合并。以本版本生成源码逐键核对：Health、AbsorptionAmount、Pos、Motion、Rotation、Fire、Air、Silent、NoAI、NoGravity、PersistenceRequired、CustomName、Tags、UUID、Passengers、Brain、Inventory、EnderItems 等仍为 PascalCase；**FallDistance → fall_distance、Attributes → attributes、Leash → leash、HomePos → home_pos、DropChances → drop_chances** 已迁移；**HandItems / ArmorItems 列表合并为 equipment map**（槽位名 mainhand / offhand / head / chest / legs / feet / body，物品形如 {id, count, components}）。SNBT 数字后缀仍为强制（{Health:10f} 有效，{Health:10} 不匹配浮点）。EntityLayer 的语义分类表（IDENTITY / NUMERIC / VALUE / SLOT / RELATIONAL）据此更新：新增 equipment→SLOT、fall_distance→NUMERIC、leash→RELATIONAL，未知名默认 VALUE 恢复的保守语义不变。", 10.5, False, "justify", 2, 6, 0)
para("1.3　目标分层", 12, True, "left", 0, 6, 6)
para("(1) **捕获层**——/data 对实体的全部路径级读/写进入 PhaseSnapshot；(2) **校验层**——NBT 读集在实体所属区域与实况比对，接通 CHECK_NBT_READ_SET；(3) **补偿层**——回滚补偿改道按实体分派的 onEntityAsync（修正 P2 的全局 tick 违规）；(4) **快路径修复**——链内同区域 /data 恢复同步语义。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 二 ----------------
para("二、总体设计", 14, True, "left", 0, 6, 12)
para("2.1　捕获点选择论证：命令层而非数据池", 12, True, "left", 0, 6, 6)
para("积分板捕获选在数据池（Scoreboard 全局结构），因为其读写汇合点唯一（getPlayerScoreInfo / 匿名 ScoreAccess.set / reset*）。实体 NBT 没有同构的汇合点：Entity.load / saveWithoutId 是全量复合标签的进出，路径解析发生在 **NbtPath.get** 之后的命令语义层。若在数据池拦截，将退化为“整实体复合标签的粗粒度读写”，丢失数值、槽位、子字段的粒度，补偿只能整实体恢复。因此 EXP6 把捕获点放在**命令层**——DataCommands 已持有 path 与内存复合快照，读旧值/读新值均可在变异前拷贝，天然获得路径粒度且不触碰实体线程外的数据。", 10.5, False, "justify", 2, 6, 0)
para("2.2　键归一：26.1 路径文法 → 稳定复合键", 12, True, "left", 0, 6, 6)
para("NbtPath.asString() 返回命令原文，文法为：简单键（Health）、向量下标（Pos[0]）、复合子字段（foo.bar）、槽位匹配（26.1 为 equipment.mainhand 的点路径；旧式 Inventory[{Slot:0b}].id 的括号形式仅存于历史存档）。**EntityLayer.parsePathString** 解析为 {tagName, slotKey, subField} 三元组，nbtKey 归一为 “{entityId}:{tagName}[/{subField}]” 或 “{entityId}:{slotKey}[/{subField}]” 稳定键；表 1 给出文法与键的对应。", 10.5, False, "justify", 2, 6, 0)
table([
    ["路径原文（26.1）", "tagName", "slotKey", "subField", "复合键"],
    ["Health", "Health", "—", "—", "123:Health"],
    ["Pos[0]", "Pos[0]", "—", "—", "123:Pos[0]"],
    ["foo.bar", "foo", "—", "bar", "123:foo/bar"],
    ["equipment.mainhand", "equipment", "—", "mainhand", "123:equipment/mainhand"],
    ["equipment.mainhand.id", "equipment", "—", "mainhand/id", "123:equipment/mainhand/id"],
])
para("2.3　线程模型：快照显式传递 + 并发容器", 12, True, "left", 0, 6, 6)
para("命令方块链内 /data 在头区域线程执行（快路径内联）或经 getDataAsync/setDataAsync 落到实体区域线程（跨区异步路径）。快照在命令入口以 **ExpChainSupport.getPhaseSnapshot()** 显式捕获并作为 final 引用传入异步回调——不复用积分板的“全局 tick 播种”机制，因为实体任务的执行线程不固定。为容忍两个执行面上对同一 Phase 的并发捕获，PhaseSnapshot 的四个 NBT 映射（nbtCache / nbtOldValues / nbtReadSet / nbtReadSetValues）改用 ConcurrentHashMap，putIfAbsent 语义保持“首旧值稳定”。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 三 ----------------
para("三、捕获接线（命令层注入）", 14, True, "left", 0, 6, 12)
para("3.1　六条执行流的捕获点", 12, True, "left", 0, 6, 6)
para("**getData / getNumeric（路径读）**：getSingleTag 解析出 Tag 后调用 interceptEntityRead → **EntityLayer.recordReadValue**：键命中本 Phase 写缓存则只计缓存命中（与积分板读钩同构，**不记读集**，避免读己写自冲突）；否则将（键, 观测值, 0）写入读集取值层。异步路径在 getDataAsync 完成回调（实体区域线程）上解析并捕获，解析后的 Tag 再交 runOnSource 回源展示。**resolveSourcePath（modify from 源路径读）**：同构捕获源实体的路径读。**manipulateData / removeData / mergeData（写）**：变异前 readLeaf 拷贝旧值，变异后 readLeaf 取新值，**写入落地后**（setData/setDataAsync 成功）调用 interceptEntityWrite——写失败不残留幻影补偿数据；remove 以 **EntityLayer.REMOVED** 哨兵标记，补偿时按“恢复旧值”处理；无路径的 merge 走顶层浅差集（after.keySet() 逐键比较）。块/存储访问器一律跳过（instanceof EntityDataAccessor 判定）。表 2 列出接线状态。", 10.5, False, "justify", 2, 6, 0)
table([
    ["执行流", "捕获内容", "执行面", "状态"],
    ["get / getNumeric（路径）",  "interceptRead 缓存判定 + recordReadValue 读集取值", "快路径内联 / 实体区域回调", "已接线"],
    ["modify from <sourcePath>", "源实体路径读集", "实体区域回调", "已接线"],
    ["modify（set/merge/insert/append/prepend）", "路径写 (new, old)", "写落地后（快/异步）", "已接线"],
    ["merge（无路径，顶层浅差集）", "逐顶层键 (new, old)", "写落地后（快/异步）", "已接线"],
    ["remove", "REMOVED 哨兵 + 旧值", "写落地后（快/异步）", "已接线"],
    ["get 无路径 / 整复合源（getSingletonSource）", "—（展示/整复合直写，MVP 不捕获）", "—", "已知限制 7.1-3"],
])
para("3.2　REMOVED 哨兵与补偿语义扩展", 12, True, "left", 0, 6, 6)
para("删除类变异没有“新值”，putNbt 的 (new, old) 二元组无法表达。引入静态哨兵 **EntityLayer.REMOVED**：interceptWrite 以 REMOVED 为新值，compensateKey 遇到 REMOVED 时直接按旧值恢复（VALUE 分支），SLOT 分支的整槽删除在写回时按槽位描述符重定位并回插（列表漂移时追加，槽位身份由槽内 Slot/槽位名保持）。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 四 ----------------
para("四、校验与补偿闭环", 14, True, "left", 0, 6, 12)
para("4.1　CHECK_NBT_READ_SET：按实体分派的实况比对", 12, True, "left", 0, 6, 6)
para("verifyReadSetAndResume 在方块读集（按区域任务）与积分板读集（全局 tick）之外新增 NBT 读集校验：nbtReadSetValues 按键前缀解析实体 ID 并分组，每组经 **RegionCommandExecutor.onEntityAsync** 在实体所属区域重读实况（readEntityNbtPath 支持向量、equipment 槽位、复合子字段三种路径形态），与观测值 Objects.equals 比对，差异或实体消失记为冲突；allOf 聚合后调用 **PhaseValidator.validate(phaseSnap, retryCount, modified, modifiedScores, modifiedNbt)** 五参重载——COMMIT / RETRY / RETRY_EXHAUSTED 与既有状态机共享。PhaseValidator 新增 CHECK_NBT_READ_SET 分支，遍历 nbtReadSet 命中外部修改标记即 RETRY。", 10.5, False, "justify", 2, 6, 0)
para("4.2　补偿改道：P2 潜在缺陷的修正", 12, True, "left", 0, 6, 6)
para("EXP5Plus 将 EntityLayer.compensate 与 ScoreLayer 一并放在 **onGlobalAsync**（全局 tick）执行——积分板是服务器全局数据，但实体 NBT 必须在其所属区域线程读写，全局线程触碰实体违反 Folia 线程契约（P2 潜在缺陷）。EXP6 拆分：积分板补偿留在全局 tick；实体补偿先以 **EntityLayer.entityIdsOf** 从 nbtKeySet 提取实体 ID，逐实体经 onEntityAsync 派发 **compensateFor**（按前缀过滤单实体键），reader/writer 在实体自身区域线程内完成 Δ 逆操作与值恢复；全部补偿 future 与方块按区域恢复、头区域重放串接。", 10.5, False, "justify", 2, 6, 0)
para("4.3　快路径修复：scheduleOrExecute 的隐蔽语义", 12, True, "left", 0, 6, 6)
para("实测发现链内 /data 的写入始终晚一 tick 生效（探针读旧值），定位到 **Folia 实体调度器 scheduleOrExecute 在实体所属 tick 线程上仍将任务排队到实体下一 tick**——EXP5 假定的“同区域 getDataAsync 同步完成 → isDone 快路径”在命令方块链内从未成立。RegionCommandExecutor.onEntityAsync 改为：**TickThread.isTickThreadFor(entity) 时内联执行**（区域线程访问本区域实体符合 Folia 线程契约），否则维持 scheduleOrExecute。修复后链内同区域 /data 恢复“命令返回即写入落地”的同步语义，读己写缓存（recordReadValue 的缓存命中分支）成为真正的兜底而非主路径。", 10.5, False, "justify", 2, 6, 0)
para("4.4　接线状态对照", 12, True, "left", 0, 6, 6)
table([
    ["组件", "设计职责", "实际接线", "状态"],
    ["读捕获：recordReadValue", "缓存透传判定 + 读集取值", "DataCommands 六流 + EntityLayer", "已接线，链实测无异常"],
    ["写捕获：interceptWrite(+REMOVED)", "putNbt(new, old) / 删除哨兵", "DataCommands 写流", "已接线"],
    ["键归一：parsePathString", "26.1 路径文法 → 稳定键", "EntityLayer 新增", "已接线"],
    ["分类字典：CATEGORY", "26.1 键名（equipment/leash/fall_distance）", "EntityLayer 更新", "已接线"],
    ["读集取值层：nbtReadSetValues", "观测值基线（并发容器）", "PhaseSnapshot 新字段+重载", "已接线"],
    ["校验：CHECK_NBT_READ_SET", "按实体区域实况比对 → RETRY", "verifyReadSetAndResume 五参调用", "已接线"],
    ["补偿：compensateFor", "按实体区域 Δ/恢复逆操作", "rollbackAndRetryExpChain 重构", "已接线"],
    ["快路径：onEntityAsync 内联", "isTickThreadFor → 同步执行", "RegionCommandExecutor 修复", "已接线"],
])

# ---------------- 五 ----------------
para("五、一致性分析", 14, True, "left", 0, 6, 12)
para("5.1　读己写（Phase 内缓存透传）", 12, True, "left", 0, 6, 6)
para("recordReadValue 先查 nbtCache：命中（本 Phase 已写该键）则只计缓存命中、不记读集——与积分板读钩同构，杜绝“读后写”在 Phase 校验时的自冲突风暴。未命中则记录观测值；快路径修复后链内同区域读发生在写落地之后，观测值即为新值，校验比对自然一致。", 10.5, False, "justify", 2, 6, 0)
para("5.2　不可重复读检测", 12, True, "left", 0, 6, 6)
para("读集取值层记录观测值，Phase 挂起恢复前在实体所属区域重读实况比对；外部写介入则置修改标记 → RETRY → 补偿 → 重放。检测窗口为“读取发生”至“Phase 恢复校验”；比对线程与实体写入线程同为实体区域线程，比较无竞态。重放期间 resetForRetry 清空读集与缓存，重放重新捕获。", 10.5, False, "justify", 2, 6, 0)
para("5.3　补偿代数与分类语义", 12, True, "left", 0, 6, 6)
para("NUMERIC（Health、Pos[0] 等）以 Δ 逆操作补偿：target = current − (new − old)，无并发精确恢复、有并发保留并发增量（Saga 语义）；VALUE（Silent、NoAI 等）与 REMOVED 按旧值恢复；SLOT（equipment 子字段/整槽）按槽位描述符重定位写回，整槽被删时回插（列表漂移时追加，槽位身份不依赖下标）；IDENTITY（id、UUID、CustomName、Tags）永不回滚；RELATIONAL（Passengers、Brain、leash）写时标记 Phase 不可逆。未知名默认 VALUE 恢复，保守安全。", 10.5, False, "justify", 2, 6, 0)
para("5.4　非保证项", 12, True, "left", 0, 6, 6)
para("与 EXP3 一脉相承：写偏斜与幻读不保证；“读取不存在”不记入读集，外部并发创建该键不被检测；路径多匹配时只捕获首个匹配（MVP）；整复合读（get 无路径）与整复合源（modify from 无 sourcePath）不捕获；实体生命周期（summon/despawn）不在补偿域。上述边界均属有意取舍，非实现缺陷。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 六 ----------------
para("六、构建与验证", 14, True, "left", 0, 6, 12)
para("6.1　构建管线", 12, True, "left", 0, 6, 6)
para("修改落于三处：azurebranches-common（PhaseSnapshot 并发 NBT 映射与读集取值层、EntityLayer 键归一/哨兵/分类字典/compensateFor、PhaseValidator 五参校验）、azurepatches-src（DataCommands 六流捕获、EntityDataAccessor.entityId）、azurepatches-new（RegionCommandExecutor.onEntityAsync 内联修复）、build.gradle.kts（verifyReadSetAndResume NBT 校验注入、rollbackAndRetryExpChain 补偿重构、readEntityNbtPath/writeEntityNbtPath 路径形态扩展）。构建日志证实全部 transformSource 锚点命中，编译零错误（仅既有 deprecation 警告）。产物 **azurebranches-server-26.1.2-AB-0002-EXP6.jar**（51 MB），构建链为 buildFolia + mergeJar。", 10.5, False, "justify", 2, 6, 0)
para("6.2　链内功能矩阵实测（EXP 模式）", 12, True, "left", 0, 6, 6)
para("环境：Windows 11 / JDK 25 / Gradle 9.4.1 / RCON 布置、服务器日志（logs/latest.log）断言。Phase A 为 15 块命令方块链（impulse + 14 chain，auto:1b），覆盖：merge Health 10f、modify Health 15、modify Pos[0] 5.5、merge equipment.mainhand 石头（26.1 装备 map 形态）、data get equipment.mainhand.id（槽位点路径读）、merge Silent、remove Silent；每步后以 execute if entity @e[...,nbt={...}] run say 探针断言（26.1 NBT 形态：PascalCase 键名 + 强制数字后缀 f/d + equipment map + 单元素列表子集匹配）。**8/8 探针全部命中**，链无异常，读捕获调试日志证实 recordReadValue 按路径粒度记录（如 26:equipment/mainhand/id → minecraft:stone、26:Health → 15.0）。", 10.5, False, "justify", 2, 6, 0)
para("6.3　OCC 读-写自冲突与重放收敛实测", 12, True, "left", 0, 6, 6)
para("Phase B：merge Silent → data get Health（读捕获 15）→ data modify Health 18（写捕获）→ 跨区 setblock（区域 (1,0)，强制挂起 + 校验）。CHECK_NBT_READ_SET 检出 live 18 ≠ 观测 15 → RETRY → 按实体区域补偿（Health Δ 回 15、Silent 恢复移除）→ 整 Phase 重放——**Phase 起始探针在日志中出现 4 轮（初始 + 3 次重试）即重放发生的铁证**；耗尽后接受提交，终态探针 Health:18f 与 Silent:1b 双双命中，服务器全程零 ERROR。该实验同时覆盖：读集捕获、写集捕获、按实体区域的实况比对、Δ/恢复补偿、方块恢复、整 Phase 重放、重试上限收敛七个子系统。", 10.5, False, "justify", 2, 6, 0)
para("6.4　验证方法学说明", 12, True, "left", 0, 6, 6)
para("Folia 异步 RCON 仅回传失败文本，成功反馈与 say 广播只进服务器日志，且控制台 stdout 管道为块缓冲——断言一律走服务器自写 logs/latest.log（log4j 实时写）；控制台源实体选择器存在跨线程竞态，实体相关命令一律经命令方块（区域线程）执行；name= 选择器在 26.1 对 CustomName 组件比较异常（getPlainTextName 语义变化），以 type= 选择器 + 清场保证唯一；EXP walker 下 impulse 头块命令不参与重放，计数探针置于链首块；time set 为异步全局命令，须先等其生效再召唤测试实体并以 Invulnerable 免疫环境伤害（白天阳光会烧伤无盔甲僵尸，造成 nbt 探针假阴性）。上述通道限制均已排除对被测语义的干扰。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 七 ----------------
para("七、已知限制与后续工作", 14, True, "left", 0, 6, 12)
para("7.1　已知限制", 12, True, "left", 0, 6, 6)
para("1. 整复合读（data get 无路径）与整复合源（modify from 无 sourcePath）不捕获——展示型读与整复合直写不产生路径粒度键。", 10.5, False, "justify", 2, 3, 0)
para("2. 路径多匹配（如通配 list 元素）仅捕获首个匹配，其余匹配的旧值不在补偿域。", 10.5, False, "justify", 2, 3, 0)
para("3. “读取不存在”不记入读集：外部并发创建该键不被 CHECK_NBT_READ_SET 捕获。", 10.5, False, "justify", 2, 3, 0)
para("4. 跨区异步路径的捕获发生在实体区域回调，与 Phase 结束校验之间存在微弱竞窗（并发容器保证安全，但极晚落地的捕获可能错过本轮校验）——快路径修复后链内同区域主场景不受影响。", 10.5, False, "justify", 2, 3, 0)
para("5. 实体生命周期（summon / despawn / 死亡）不在补偿域；SUMMON 延迟动作仍未接线（延续 EXP5 遗留）。", 10.5, False, "justify", 2, 3, 0)
para("7.2　后续工作（优先级排序）", 12, True, "left", 0, 6, 6)
para("**P0**：SUMMON 延迟动作接线（DeferredAction 队列的最后一个缺口）；/team 恢复（复用全局 tick 范式）。", 10.5, False, "justify", 2, 3, 0)
para("**P1**：捕获计数器观测出口（nbtInterceptWrite / nbtCacheHit / nbtCompensationCount 暴露为可查询统计）；双链并发写同实体 NBT 的确定性冲突编排。", 10.5, False, "justify", 2, 3, 0)
para("**P2**：整复合读/写的路径粒度展开（对复合读做顶层键分解捕获）；跨区异步路径的捕获时序收紧（将实体区域捕获 future 纳入 EXP 链回执袋）。", 10.5, False, "justify", 2, 3, 0)

# ---------------- 八 ----------------
para("八、结论", 14, True, "left", 0, 6, 12)
para("EXP6 的意义在于三点。其一，EntityLayer 从骨架变为闭环：捕获点选择论证（命令层 vs 数据池）、26.1 键名迁移的逐键核对、路径文法到稳定复合键的归一、按实体区域的实况校验与 Δ/恢复补偿，使实体 NBT 成为第三个完整走通“捕获 → 验证 → 补偿”的 OCC 数据域。其二，修正了 EXP5Plus 的 P2 潜在缺陷——实体补偿从全局 tick 改道按实体分派，线程契约恢复。其三，工程上挖出 Folia 实体调度器 scheduleOrExecute 的隐蔽语义陷阱并以内联执行修复，链内 /data 的同步快路径首次真正生效——这是比功能接线更底层的正确性收获。", 10.5, False, "justify", 2, 6, 0)
para("方法学上，本版本再次验证：以构建期锚点唯一性校验对冲上游漂移、以生成源码逐键核对对抗版本迁移、以日志广播通道在异步 RCON 的受限反馈下完成确定性断言。三者结合，使 Folia 的多线程模型与命令方块的确定性语义在实体数据域上相容。", 10.5, False, "justify", 2, 6, 0)
para("文档：AzureBranches-26.1.2-EXP6　　版本：26.1.2-EXP6　　日期：2026 年 08 月 16 日　　项目：AzureBranches (https://github.com/XCxyTianQ/AzureBranches)", 9, False, "center", 0, 12, 0)

doc.save(OUT)
print("DOCX written:", OUT)
