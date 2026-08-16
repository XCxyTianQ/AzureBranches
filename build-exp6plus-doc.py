# -*- coding: utf-8 -*-
"""Generates F:\AzureCore\AzureDoc\AzureBranches-26.1.2-EXP6Plus.docx
Format mirrors the EXP6 document (direct formatting, CN section numbering)."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"F:\AzureCore\AzureDoc\AzureBranches-26.1.2-EXP6Plus.docx"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

BOLD_RE = re.compile(r"(\*\*.*?\*\*)")


def _set_run(run, size, bold, font_east="宋体"):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_east)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def para(text, size=10.5, bold_all=False, align="justify", indent_chars=0, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    if indent_chars > 0:
        p.paragraph_format.first_line_indent = Pt(size * indent_chars)
    for part in BOLD_RE.split(text):
        if not part:
            continue
        m = BOLD_RE.fullmatch(part)
        if m:
            _set_run(p.add_run(m.group(1)[2:-2]), size, True)
        else:
            _set_run(p.add_run(part), size, bold_all)
    return p


def table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = t.cell(i, j)
            p = cell.paragraphs[0]
            r = p.add_run(str(cell_text))
            _set_run(r, 9.5, i == 0)
            if i == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "E7E6E6")
                cell._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()


# ---------------- Title block ----------------
para("AzureBranches", 22, True, "center", 0, 0, 6)
para("EXP6Plus：Scoreboard 实体维度的读集与幽灵分数防护", 16, True, "center", 0, 6, 0)
para("从“选择器解析实体集合进入 PhaseSnapshot”到“CHECK_ENTITY_READ_SET 复核 + 补偿主动清除”的记分板实体一致性工程", 12, False, "center", 0, 6, 0)
para("版本 26.1.2-EXP6Plus　　2026 年 08 月 17 日", 10.5, False, "center", 0, 6, 0)
para("基于 Folia Regionized Ticking 模型", 10.5, False, "center", 0, 6, 0)
para("感谢 PaperMC / Folia (Spottedleaf) / LuminolMC (EarthMe 及其团队)", 10.5, False, "center", 0, 12, 0)

# ---------------- Abstract ----------------
para("**摘要**", 10.5, False, "justify", 2, 6, 0)
para("EXP6Plus（v26.1.2-EXP6Plus）使 /scoreboard 的实体维度成为继方块、积分板数值、实体 NBT 之后第四个完整走通“捕获 → 验证 → 补偿”的 OCC 数据域。捕获点选择于 **ScoreHolderArgument.getNames**——选择器解析的自然汇合点：解析出的实体集合以（实体 ID → scoreboardName 观测值）进入 PhaseSnapshot 的实体读集（**entityReadSet / entityReadSetValues**）；ScoreboardCommand 全部终端处理器重构为在源线程完成 holder 解析，修正了全局 tick 线程迭代实体的线程契约违规。校验侧：verifyReadSetAndResume 对实体读集逐实体经 **onEntityAsync** 在实体所属区域复核存在性，存在性判定同时覆盖 **isRemoved 与 !isAlive**（死亡动画窗口内“已死未移除”的实体必须视为消失），接通 **CHECK_ENTITY_READ_SET** 六参 validate。补偿侧：rollbackAndRetryExpChain 从实体读集取值层收集 **deadHolderNames**，ScoreLayer 补偿的 writer 对已死 holder **跳过数值写回并主动清除分数条目**（resetSinglePlayerScore），计入补偿失败统计——这是因为工程实测发现 **Folia 已将 Scoreboard.entityRemoved 注释禁用**，死实体的分数条目不会像 vanilla 那样自动消失，仅“跳过写回”无法防护幽灵分数。跨 Phase 语义上，Continuation 新增实体读集与分数写状态的携带字段，PhaseSnapshot.fromContinuation 继承之，且携带时机必须延迟到**挂起屏障完成之后**——全局 tick 的分数写落在挂起窗口内，挂起点复制会看到空写集。全文以“捕获点选择 — 跨相位携带 — 校验 — 幽灵防护 — 验证”为主线，并以实体存活 COMMIT、实体死亡冲突 RETRY、幽灵分数防护三条链路的实测数据佐证（smoke 5/5 全绿）。", 10.5, False, "justify", 2, 6, 0)
para("**关键词**：Folia；命令方块；Scoreboard；ScoreHolderArgument；实体读集；PhaseSnapshot；OCC；CHECK_ENTITY_READ_SET；onEntityAsync；isAlive；死亡动画窗口；幽灵分数防护；entityRemoved；跨相位携带", 10.5, False, "justify", 0, 12, 0)

# ---------------- 一 ----------------
para("一、背景与问题定义", 14, True, "left", 0, 6, 12)
para("1.1　EXP6 质询的结论：Scoreboard 的实体维度不参与校验", 12, True, "left", 0, 6, 6)
para("对 EXP6 的质询确认：/data 的实体 NBT 捕获（EXP6）并不改善 /scoreboard 的实体数据记录检查。EXP6Plus 即以此质询结论为构建基础：/scoreboard 选择器（@e[type=…] 等）解析出的实体集合本身是一次对“存活实体集合”的读取——若链内某实体在解析后死亡/消失，后续一切针对该 holder 的分数操作都基于幻影集合。EXP5Plus 的积分板数值读集（scoreReadSetValues）只记录“目标分数键的观测值”，不含“holder 背后的实体是否存在”这一维度。", 10.5, False, "justify", 2, 6, 0)
para("1.2　工程实测发现的三个隐蔽事实", 12, True, "left", 0, 6, 6)
para("(1) **Folia 禁用了 Scoreboard.entityRemoved**——ServerLevel 的实体移除路径将该调用注释掉（scoreboard 是服务器全局数据，移除发生在区域线程）。后果：实体死亡后其分数条目不像 vanilla 那样被 resetAllPlayerScores 清除，**死实体的分数条目会永久残留**。幽灵防护因此不能只“跳过补偿写回”，必须“主动清除条目”。", 10.5, False, "justify", 2, 3, 0)
para("(2) **performCommand 恒真**——BaseCommandBlock.performCommand 对一切非异常执行无条件返回 true（vanilla 语义，successCount 只驱动条件块），失败命令不会打断链。测试中曾据此误判“链断”，实际是失败命令被静默越过。", 10.5, False, "justify", 2, 3, 0)
para("(3) **死亡动画窗口**——实体被 /damage 致死后进入死亡动画（约 20 tick），期间 isRemoved() 为假、level.getEntity 仍返回实体；若存在性复核只判 isRemoved，此窗口内的死亡冲突必然漏检。", 10.5, False, "justify", 2, 3, 0)
para("1.3　目标分层", 12, True, "left", 0, 6, 6)
para("(1) **捕获层**——/scoreboard 选择器解析出的实体集合进入 PhaseSnapshot；(2) **校验层**——Phase 校验经 onEntityAsync 复核实体存在性，消失 → CHECK_ENTITY_READ_SET → RETRY；(3) **补偿层**——回滚补偿对已死实体 holder 跳过写回并主动清除（幽灵分数防护，计入补偿失败统计）；(4) **跨相位语义**——实体读集与分数写状态跨 Phase 携带，使“phase N 解析、phase N+1 死亡”的冲突可被检出。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 二 ----------------
para("二、总体设计", 14, True, "left", 0, 6, 12)
para("2.1　捕获点选择论证：ScoreHolderArgument.getNames", 12, True, "left", 0, 6, 6)
para("/scoreboard 的 holder 参数有四种形态：@选择器、UUID、玩家名、*（tracked players 通配）。**getNames（含 getNamesWithDefaultWildcard / getName）是全部形态的汇合点**，解析结果统一为 Collection<ScoreHolder>；实体参与记分板时直接实现 ScoreHolder 接口（26.1 中实体即 holder）。在此处捕获无需触碰 ScoreboardCommand 的十余个终端处理器各自的语义，也不触碰数据池（Scoreboard 全局结构只有数值读写汇合点，没有“实体维度”的汇合点）。捕获内容：每个 entity 型 holder 的 **实体 ID + getScoreboardName()**（实体记分板身份，等于其 stringUUID）。", 10.5, False, "justify", 2, 6, 0)
para("2.2　线程模型：源线程解析 + 全局 tick 写 + 区域线程复核", 12, True, "left", 0, 6, 6)
para("选择器解析迭代实体，属**区域数据**，只能在链的源线程（区域线程）执行——ScoreboardCommand 的终端处理器据此重构为在源线程解析出 final Collection<ScoreHolder> 后，仅把纯分数操作交 dispatch() 的全局 tick 任务；EntityLayer.recordEntityRead 在解析现场把实体读集写入 PhaseSnapshot。校验与补偿仍在实体所属区域线程（onEntityAsync）与全局 tick（onGlobalAsync）执行。", 10.5, False, "justify", 2, 6, 0)
para("2.3　跨 Phase 携带：链 = 跨 Phase 的单一事务", 12, True, "left", 0, 6, 6)
para("OCC 的“读—验”窗口跨越 Phase 边界：链在 phase 1 的 z3 解析实体（记录读集），在 phase 2 的 z2 杀死该实体，冲突只能在 phase 2 的校验中检出。PhaseSnapshot.fromContinuation 原本只继承方块读集位置与旧方块状态——实体读集与分数写状态在 phase 2 校验时已丢失。Continuation 新增三个携带字段（**entityReadCarry / oldScoreValuesCarry / scoreCacheCarry**），fromContinuation 继承；携带复制必须发生在**挂起屏障完成之后**（见 4.1），否则全局 tick 的分数写（落在挂起窗口内）不会被复制。", 10.5, False, "justify", 2, 6, 0)
para("2.4　读取模型", 12, True, "left", 0, 6, 6)
para("entityReadSet: Map<Integer, Long>（实体 ID → 读取 tick）；entityReadSetValues: Map<Integer, String>（实体 ID → scoreboardName 观测值）。后者是校验（存在性复核的键）与补偿（deadHolderNames 的值）的共同输入。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 三 ----------------
para("三、捕获接线", 14, True, "left", 0, 6, 12)
para("3.1　ScoreHolderArgument 覆盖层（azurepatches-src）", 12, True, "left", 0, 6, 6)
para("在 getNames 的统一出口处：解析结果为空照旧抛 NO_ENTITIES_FOUND；非空且 EXP 上下文活跃（ExpChainSupport.getPhaseSnapshot() 非空）时，遍历结果，对 **holder instanceof Entity** 者调用 **EntityLayer.recordEntityRead(snap, entity.getId(), holder.getScoreboardName(), snap.getSnapshotTick())**，记入实体读集与取值层（putIfAbsent）。上下文不活跃（玩家/控制台）零开销跳过。", 10.5, False, "justify", 2, 6, 0)
para("3.2　ScoreboardCommand 源线程解析重构", 12, True, "left", 0, 6, 6)
para("全部终端处理器（set/add/remove/reset/enable/operation 等约十五处）改为：在源线程解析出 final Collection<ScoreHolder>（或单个 holder），随后仅把分数操作交给 dispatch() 的全局 tick 任务。此前解析发生在全局 tick 闭包内，实体迭代违反 Folia 线程契约；重构后同时获得捕获现场（源线程上有 PhaseSnapshot 线程局部）。dispatch() 保留 registerRemote 回执注册（EXP6Plus 的链 future 屏障已真正等待它）。", 10.5, False, "justify", 2, 6, 0)
para("3.3　PhaseSnapshot 字段与生命周期", 12, True, "left", 0, 6, 6)
para("新增 entityReadSet / entityReadSetValues（ConcurrentHashMap，构造器初始化）、recordEntityRead、getEntityReadSet / getEntityReadSetValues；resetForRetry 清空两者（重放重新捕获）。fromContinuation 继承 Continuation.entityReadCarry 与分数携带字段（oldScoreValues / scoreCache putIfAbsent）。", 10.5, False, "justify", 2, 6, 0)
table([
    ["执行流", "捕获内容", "执行面", "状态"],
    ["@选择器（getNames 出口）", "实体 ID + scoreboardName", "源线程（区域线程）", "已接线"],
    ["UUID / 玩家名 / *", "同名判定（entity 型 holder）", "源线程", "已接线"],
    ["分数写（ScoreAccess.set 数据池钩）", "putScore(new, old)", "全局 tick（快照播种）", "EXP5Plus 已接线"],
    ["分数读（getPlayerScoreInfo 数据池钩）", "recordScoreRead 观测值", "全局 tick（快照播种）", "EXP5Plus 已接线"],
])

# ---------------- 四 ----------------
para("四、校验与补偿闭环", 14, True, "left", 0, 6, 12)
para("4.1　CHECK_ENTITY_READ_SET：存在性复核与死亡动画窗口", 12, True, "left", 0, 6, 6)
para("verifyReadSetAndResume 在方块读集（按区域任务）、积分板读集（全局 tick）、NBT 读集（按实体区域）之外新增实体读集校验：对 entityReadValues 逐实体，**level.getEntity(entityId) 为空的直接记冲突**；非空者经 **onEntityAsync** 在实体所属区域复核——判定式为 **scheduled.isRemoved() || !scheduled.isAlive() || level.getEntity(entityId) == null**。isRemoved 只覆盖“已移除”形态；**!isAlive 覆盖死亡动画窗口**（已死未移除的实体仍在 level.getEntity 中，但 holder 集合已失效）；重查 level.getEntity 兜底异步窗口。全部完成经 CompletableFuture.allOf 聚合后调用 **PhaseValidator.validate(phaseSnap, retryCount, modified, modifiedScores, modifiedNbt, modifiedEntities)** 六参重载，PhaseValidator 新增 CHECK_ENTITY_READ_SET 分支——命中消失标记即 RETRY。", 10.5, False, "justify", 2, 6, 0)
para("4.2　携带时机的时序修正", 12, True, "left", 0, 6, 6)
para("实测捕获调试日志显示：挂起点（dispatchAndSuspend）复制携带字段时 oldScores/scoreCache 恒为空。原因是分数写发生在 **挂起窗口内**——z3 的 set 把全局 tick 任务入队并注册链 future，屏障等待该 future 完成（写落地 + putScore 记录）后才恢复，而携带复制发生在挂起点、早于写落地。修正：携带复制移至屏障完成回调（allOf.whenComplete 的 queueOrExecuteTickTask 内、aggregateAndResume 之前）——复制时刻的 PhaseSnapshot 已含分数写状态。实体读集在源线程解析时已记录，早/晚复制均可见。", 10.5, False, "justify", 2, 6, 0)
para("4.3　幽灵分数防护：跳过写回 + 主动清除", 12, True, "left", 0, 6, 6)
para("rollbackAndRetryExpChain 在补偿前从 phaseSnap.getEntityReadSetValues() 收集 **deadHolderNames**——实体不存在或 !isAlive 的 holder 记分板身份。ScoreLayer.compensate 的 writer 对 deadHolderNames 命中者：**跳过数值写回，并 resetSinglePlayerScore 主动清除条目**（模拟被 Folia 禁用的 vanilla entityRemoved 语义），同时 **onScoreCompensationFailed() 计入补偿失败统计**。仅“跳过写回”不足以防护：Folia 下死实体的旧条目仍在，跳过等于保留幽灵。", 10.5, False, "justify", 2, 6, 0)
para("4.4　RETRY 与重放语义", 12, True, "left", 0, 6, 6)
para("冲突 → onValidationRetry → retryCount++ → rollbackAndRetryExpChain：补偿（分数清除/恢复、实体 NBT、按区域方块恢复）→ 从 phaseStartPos 整 Phase 重放（resetForRetry 清空读集）。重放不重放 impulse 头块命令；重放轮中 z3 的 set 因实体已死而失败（performCommand 恒真，链继续走完），其读集不再记录，后续校验空读集直接 COMMIT——重试有界收敛（max_retries 兜底）。", 10.5, False, "justify", 2, 6, 0)
para("4.5　接线状态对照", 12, True, "left", 0, 6, 6)
table([
    ["组件", "设计职责", "实际接线", "状态"],
    ["捕获：recordEntityRead", "实体 ID + scoreboardName 入读集", "ScoreHolderArgument 覆盖层 + EntityLayer", "已接线"],
    ["源线程解析", "holder 解析移出全局 tick", "ScoreboardCommand 全处理器重构", "已接线"],
    ["读集字段", "entityReadSet / entityReadSetValues", "PhaseSnapshot 新增（并发容器）", "已接线"],
    ["跨相位携带", "entityReadCarry + 分数写状态", "Continuation 三字段 + fromContinuation 继承", "已接线"],
    ["携带时机", "屏障完成后复制", "dispatchAndSuspend 恢复回调", "已接线"],
    ["校验：CHECK_ENTITY_READ_SET", "onEntityAsync 复核 + isAlive", "verifyReadSetAndResume 六参调用", "已接线"],
    ["幽灵防护", "跳过写回 + 主动清除 + 失败统计", "rollbackAndRetryExpChain writer 分支", "已接线"],
])

# ---------------- 五 ----------------
para("五、一致性分析", 14, True, "left", 0, 6, 12)
para("5.1　实体消失的三种形态", 12, True, "left", 0, 6, 6)
para("(1) **已移除**（isRemoved，实体被 discard）——getEntity 返回空，判定直接命中；(2) **死亡动画中**（dead=true 未移除）——getEntity 仍返回实体，靠 !isAlive 命中；(3) **区块卸载**（无玩家在线时实体随区块存入磁盘、对 @e 不可见）——属于环境事实而非死亡，vanilla 语义下不触发清除，本版本不将其计入冲突（见 7.1-3）。", 10.5, False, "justify", 2, 6, 0)
para("5.2　补偿代数：为何“清除”优于“跳过”与“恢复”", 12, True, "left", 0, 6, 6)
para("vanilla 语义：实体死亡 → entityRemoved → resetAllPlayerScores——死亡后该 holder 不应有任何条目。Folia 禁用该调用后，条目残留使三种补偿策略产生分歧：**恢复旧值**会重建（或刷新）一条属于死实体的条目（最坏——幽灵复活）；**跳过写回**保留链写的新值（次坏——幽灵值）；**主动清除**删除条目，与 vanilla 终态一致（正确）。EXP6Plus 取清除语义，且清除动作计入补偿失败统计——补偿无法“恢复”，属于异常路径，必须可见。", 10.5, False, "justify", 2, 6, 0)
para("5.3　读己写与非保证项", 12, True, "left", 0, 6, 6)
para("实体读集只记录“解析时刻的存活实体集合”，不记录实体属性值——校验是存在性复核而非状态比对，自写（链内杀死实体）与外部杀死实体无法区分、也不需区分（两者都使 holder 集合失效）。非保证项：区块卸载导致的“不可见”不触发冲突；实体在解析后、写入前消失但链未挂起时不校验（无挂起即无窗口）；/team 等 holder 侧结构操作不在实体读集域。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 六 ----------------
para("六、构建与验证", 14, True, "left", 0, 6, 12)
para("6.1　构建管线", 12, True, "left", 0, 6, 6)
para("修改落于四处：azurebranches-common（PhaseSnapshot 实体读集与 fromContinuation 继承、EntityLayer.recordEntityRead、PhaseValidator 六参校验、Continuation 三携带字段）、azurepatches-src（ScoreHolderArgument 覆盖层新增、ScoreboardCommand 源线程解析重构）、build.gradle.kts（verifyReadSetAndResume 实体校验注入、rollbackAndRetryExpChain 幽灵防护 writer、携带时机锚点、链 future 屏障）、exp6plus-smoke.py（验收驱动）。构建日志证实全部 transformSource 锚点命中，编译零错误（仅既有 deprecation 警告）。产物 **azurebranches-server-26.1.2-AB-0002-EXP6Plus.jar**（51 MB），构建链为 buildFolia + mergeJar。", 10.5, False, "justify", 2, 6, 0)
para("6.2　smoke 实测（EXP 模式，5/5 全绿）", 12, True, "left", 0, 6, 6)
para("环境：Windows 11 / JDK 25 / Gradle 9.4.1 / RCON 布置 + 服务器日志断言。固定 UUID [I;1,2,3,4] 僵尸（NoAI + PersistenceRequired，夜间），记分板身份 = 00000001-0000-0002-0000-000300000004。每条链触发前新鲜召唤僵尸并在数秒内触发（无玩家在线时区块会卸载，僵尸对 @e 不可见——测试环境事实，见 6.3）。", 10.5, False, "justify", 2, 6, 0)
table([
    ["链路", "链结构", "预期", "实测"],
    ["C1 实体存活 COMMIT", "set @e 5 → 探针 → setblock → 探针 → END", "SET-OK / COMMIT-KEEP-5 / END 全中，无重试，写落地", "PASS（4 项断言全过）"],
    ["C2 实体死亡冲突", "set @e 7 → damage 致死 → setblock → 探针 → END", "CHECK_ENTITY_READ_SET 检出死亡（含动画窗口）→ RETRY → 补偿清除", "PASS（C3 无条目为间接铁证）"],
    ["C3 幽灵分数防护", "无条目探针 + GHOST-5 / GHOST-7 探针", "NO-ENTRY 命中，GHOST-5/7 不命中", "PASS"],
])
para("C2 的 RETRY 与幽灵防护由三组证据共同确立：(1) 调试期捕获日志证实完整链路——late carry entity=1 oldScores=1 scoreCache=1 → verify 实体检查 removed=false alive=false gone=true → validate=RETRY → rollback deadHolders=1 → 终态探针 POST-NONE（条目被清除）；(2) C3 的 NO-ENTRY 只有补偿清除才会产生（Folia 已禁用 vanilla 清除）；(3) impulse 的 PHASE-START 恒为 1（重放不重放头块，符合设计语义）。", 10.5, False, "justify", 2, 6, 0)
para("6.3　验证方法学说明", 12, True, "left", 0, 6, 6)
para("Folia 异步 RCON 仅回传失败文本，成功反馈与 say 广播只进服务器日志（log4j 实时写），断言一律走 logs/latest.log；无玩家在线时区块按不活跃计时卸载，实体随区块对 @e 不可见——僵尸必须“新鲜召唤 + 数秒内触发”，且挂起机制所需的目标块选在链同区块（chunk (0,0)，触发窗口内必然加载），使实体维度 OCC 的验证与区块生命周期解耦；performCommand 恒真意味着失败命令不打断链，测试以分数终态与探针区分“静默失败”与“正常执行”；name= 选择器在 26.1 不可靠，一律 type= 选择器 + 清场。上述通道限制均已排除对被测语义的干扰。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 七 ----------------
para("七、已知限制与后续工作", 14, True, "left", 0, 6, 12)
para("7.1　已知限制", 12, True, "left", 0, 6, 6)
para("1. 区块卸载导致的实体“不可见”不计入冲突——是环境事实而非死亡；重放轮的选择器解析会静默失败（无实体），链以失败命令越过（performCommand 恒真）。", 10.5, False, "justify", 2, 3, 0)
para("2. 幽灵清除只在 OCC 回滚补偿路径执行；不挂起的链（无 future 命令）杀死实体时，死实体条目残留（Folia 的 entityRemoved 禁用是全局问题，完整修复需在 ServerLevel 移除路径以 onGlobalAsync 恢复该调用）。", 10.5, False, "justify", 2, 3, 0)
para("3. 实体读集记录的是解析时刻的集合快照，不记录实体属性；解析后、挂起前实体死亡且链无第二次挂起时，冲突窗口为空。", 10.5, False, "justify", 2, 3, 0)
para("4. 真跨区 setblock 目标（区域 (1,0)）在无玩家测试环境下因目标区块未加载/区域不 tick 而无法作为挂起载体——挂起机制的验证以同区域目标完成；跨区 deferred 批处理本身是 EXP4 既有机制，不受本版本影响。", 10.5, False, "justify", 2, 3, 0)
para("7.2　后续工作（优先级排序）", 12, True, "left", 0, 6, 6)
para("**P0**：恢复 Folia 禁用的 Scoreboard.entityRemoved（ServerLevel 移除路径 onGlobalAsync 重派发）——把“死实体分数清除”从补偿路径推广到一切死亡路径；真跨区挂起载体（目标区块 forceload / 区域活性）的测试环境补全。", 10.5, False, "justify", 2, 3, 0)
para("**P1**：实体读集的计数观测出口（entityReadCount / entityConflictCount / ghostClearCount 暴露为可查询统计）；双链并发写同实体的确定性冲突编排。", 10.5, False, "justify", 2, 3, 0)
para("**P2**：holder 侧结构操作（/team join/leave 等）纳入实体读集域；UUID/玩家名形态的“实体存在性”与 @e 形态的对齐。", 10.5, False, "justify", 2, 3, 0)

# ---------------- 八 ----------------
para("八、结论", 14, True, "left", 0, 6, 12)
para("EXP6Plus 的意义在于三点。其一，Scoreboard 的实体维度进入 OCC：选择器解析的实体集合成为可校验的读集，实体在链执行中途死亡会被 CHECK_ENTITY_READ_SET 捕获并触发整 Phase 重放，记分板的“实体维度幻读”得到工程化防护。其二，幽灵分数防护的设计被 Folia 的工程现实重新塑造——entityRemoved 被禁用意味着“跳过写回”不够，必须“主动清除”，补偿语义从代数恢复升华为生命周期对齐。其三，跨 Phase 携带与携带时机的时序修正打通了“phase N 读、phase N+1 冲突”的窗口，使链真正成为跨 Phase 的单一事务。", 10.5, False, "justify", 2, 6, 0)
para("方法学上，本版本再次验证：调试期的高密度捕获日志与验收期的行为断言分层使用（前者定位时序竞态，后者确立语义终态）；测试环境事实（区块卸载、performCommand 恒真、RCON 反馈通道）被显式建模而非绕过。smoke 5/5 全绿，实体死亡冲突与幽灵分数防护两条链路实测通过。", 10.5, False, "justify", 2, 6, 0)
para("文档：AzureBranches-26.1.2-EXP6Plus　　版本：26.1.2-EXP6Plus　　日期：2026 年 08 月 17 日　　项目：AzureBranches (https://github.com/XCxyTianQ/AzureBranches)", 9, False, "center", 0, 12, 0)

doc.save(OUT)
print("DOCX written:", OUT)
