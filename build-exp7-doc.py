# -*- coding: utf-8 -*-
"""Generates F:\AzureCore\AzureDoc\AzureBranches-26.1.2-EXP7.docx
EXP7: b_linear 存储引擎移植 + Region File 格式 v4（策略 B）。"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"F:\AzureCore\AzureDoc\AzureBranches-26.1.2-EXP7.docx"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

BOLD_RE = re.compile(r"(\*\*.*?\*\*)")


def _set_run(run, size, bold, font_east="宋体"):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_east)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def para(text, size=10.5, bold_all=False, align="justify", indent_chars=0, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    if indent_chars > 0:
        p.paragraph_format.first_line_indent = Pt(size * indent_chars)
    for part in BOLD_RE.split(text):
        if not part:
            continue
        m = BOLD_RE.fullmatch(part)
        if m:
            _set_run(p.add_run(m.group(1)[2:-2]), size, True)
        else:
            _set_run(p.add_run(part), size, bold_all)
    return p


def table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = t.cell(i, j)
            p = cell.paragraphs[0]
            r = p.add_run(str(cell_text))
            _set_run(r, 9.5, i == 0)
            if i == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "E7E6E6")
                cell._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()


# ---------------- Title block ----------------
para("AzureBranches", 22, True, "center", 0, 0, 6)
para("EXP7：b_linear 存储引擎整合与 Region File 格式 v4", 16, True, "center", 0, 6, 0)
para("从“Luminol/Arbor b_linear 移植”到“读兼容 v0x02/0x03 + 写 VERSION 4 + 配置开关（默认 MCA）”的存储引擎工程", 12, False, "center", 0, 6, 0)
para("版本 26.1.2-EXP7　　2026 年 08 月 22 日", 10.5, False, "center", 0, 6, 0)
para("基于 Folia Regionized Ticking 模型", 10.5, False, "center", 0, 6, 0)
para("感谢 PaperMC / Folia (Spottedleaf) / LuminolMC (EarthMe 及其团队) / Arbor (Little 及 xymb 血统)", 10.5, False, "center", 0, 12, 0)

# ---------------- Abstract ----------------
para("**摘要**", 10.5, False, "justify", 2, 6, 0)
para("EXP7（v26.1.2-EXP7，策略 B）把 Luminol/Arbor 的 b_linear 存储体系移植进 AzureBranches 构建管道，并以**新格式 VERSION 4** 改进之：读兼容上游 v0x02（整文件 ZSTD）/v0x03（桶式布局），写侧一律产出 v4。移植面（azartepatches-new + azurepatches-src + transformSource 锚点）覆盖 **IRegionFile 抽象、BufferedLinearRegionFile 主引擎（swap 工作文件 + master 快照 + 每桶写纪元 + 60%/1MiB 自动压缩）、BufferedLinearRegionFileFlusher、RegionFormat 格式选择器与 RegionFormatBootstrap 配置引导**，并以同一包内的 RegionFileAdapter 公共化 vanilla RegionFile 受保护方法；RegionFileStorage 覆盖层把类型从 RegionFile 整体迁移为 IRegionFile，创建点收敛为 `RegionFormat.open(...)`，Moonrise 侧的 ChunkSystemRegionFileStorage / MoonriseRegionFileIO / ChunkSystemChunkBuffer / RegionFile 四处签名锚点随之迁移。v4 相对上游的改进（F1–F8）：**每桶 16B 位置表项（offset+compressedLen+xxhash32）**、**footer 位置表 xxhash64**、**加载时边界+长度+桶哈希分层校验（任一层失败即 IOException 且携带文件与偏移，绝不返回部分数据）**、**未改动桶整块搬运（v3→v4 免重压缩）**、**加载失败后拒绝以空状态覆写 master（防空写损坏）**、**64 位时间戳沿用**、**读兼容矩阵**、**每文件独立保存调度与背压（flusher）**。验证分两层：JVM 级 harness（round-trip 47 chunk 写-关-开-读全对、桶数据/位置表/截断三类损坏注入全部检测、写中途硬杀后已提交 chunk 完好、**真实上游引擎生成的 v0x03 文件全部读回并原位迁移为 v4 后再读回**、64×64KiB 微基准 44.8 vs 12.5 MiB/s 约 3.6× 优于 MCA）；服务端级冒烟（全新存档 + `storage.region_format=b_linear_v4` 启动、三维度 + entities 共 14 个 v4 主文件、优雅 stop 后重启、硬 kill 后重启全部无错）。默认仍为 MCA，`storage.region_format` / `storage.linear.compression_level` 双配置切换。", 10.5, False, "justify", 2, 6, 0)
para("**关键词**：Folia；RegionFile；b_linear；BufferedLinearRegionFile；swap；XXHash32；ZZST；zzstd-jni；零分配哈希；IRegionFile；RegionFormat；VERSION 4；懒加载；损坏检测；崩溃一致性；迁移器", 10.5, False, "justify", 0, 12, 0)

# ---------------- 一 ----------------
para("一、背景与问题定义", 14, True, "left", 0, 6, 12)
para("1.1　原生 MCA 写入路径的局限", 12, True, "left", 0, 6, 6)
para("vanilla RegionFile 的每 chunk 写入都要重写整个 8KiB 定位表（12 字节/槽 × 1024 槽），整文件仅一个 xxhash 摘要且只覆盖头部；chunk 数据无独立校验；Oversized 需要额外的 .oversized 附属文件；写失败路径存在静默。多区域并行写入时，定位表重写是原生模式的主要写放大来源。", 10.5, False, "justify", 2, 6, 0)
para("1.2　方案对比与策略选择（策略 B）", 12, True, "left", 0, 6, 6)
para("候选 A（仅替换 MCA 文件格式）保留 RegionFile 结构语义，但默认路径的写放大与损坏面不变；**策略 B**（用户选定）完整移植 Luminol/Arbor 的 b_linear 引擎：swap 工作文件（LZ4 段、1024 Sector、DELETE_ON_CLOSE）承载所有写入，master 快照仅在回写窗口（写后超时 3s / 启用后首写 / 关闭）同步；桶（16×64 chunk）惰性加载避免启动时全量读取；每桶写纪元（epoch）驱动脏检查。在此基础上 EXP7 增加**格式版本 v4**：读兼容上游桶式/整文件布局，写一律 v4，并把校验错位与长度失配/截断等损坏形态全部提前为显式 IOException。", 10.5, False, "justify", 2, 6, 0)
para("1.3　目标分层", 12, True, "left", 0, 6, 6)
para("(1) **集成层**——引擎作为 azurepatches-new 类编译进 folia 构建，RegionFileStorage 覆盖层与 Moonrise 锚点迁移把创建点收敛到单一入口；(2) **格式层**——v4 布局 + 兼容读取矩阵 + 校验链；(3) **配置层**——默认 MCA，`storage.region_format` 选择 b_linear_v4；(4) **验证层**——JVM harness（round-trip/损坏/杀灭/迁移/性能）+ 服务端冒烟（新世界启动/停止/重启/硬杀恢复）。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 二 ----------------
para("二、总体设计", 14, True, "left", 0, 6, 12)
para("2.1　抽象与选择器", 12, True, "left", 0, 6, 6)
para("`IRegionFile extends ChunkSystemRegionFile, AutoCloseable` 暴露 getPath/getChunkDataInputStream/doesChunkExist/getChunkDataOutputStream/flush/clear/hasChunk/close/write/getOversizedData/isOversized/recalculateHeader/setOversized（含 moonrise$ 接口方法与默认 getRecalculateCount）。`RegionFormat` 枚举（MCA/B_LINEAR_V4）是唯一创建点：`RegionFormat.open(info, path, folder, sync)` 在每次区域文件创建时先经 `RegionFormatBootstrap.ensureApplied()` 应用配置（幂等、配置未加载自动重试、非法值一次性告警回退 MCA），再委托当前枚举的 create。MCA 分支由同包 `RegionFileAdapter`（extends RegionFile，公共化 write/oversized/recalculate 系列）适配，保证默认路径零行为漂移。", 10.5, False, "justify", 2, 6, 0)
para("2.2　v4 主文件布局", 12, True, "left", 0, 6, 6)
para("superblock=-0x200812250269L：header 14B（superblock | version=0x04 | compressionLevel | xxHash32Seed）+ 位置表 256B（16 × 16B：bucketOffset(8) | compressedLen(4) | bucketXXHash32(4)，全 0=无数据）+ 桶数据区（每个脏桶 rawLen(4) | compressedLen(4) | ZSTD）+ footer 21B（superblock | version | positionTableXXHash64(8) | reserved(4)）。桶哈希覆盖**整个桶块**（rawLen+compressedLen+ZSTD）；表哈希覆盖整张 256B 位置表；未改动桶按 v3/v4 同构编码整块搬运（免重压缩）。", 10.5, False, "justify", 2, 6, 0)
para("2.3　交换文件（swap）与写入模型", 12, True, "left", 0, 6, 6)
para("swap（r.<rx>.<rz>.mca.swp）沿用上游：超块 0x1145141919810L、版本 0x02、header＋1024×Sector(17B)、LZ4 段、60%/1MiB 自动压缩、DELETE_ON_CLOSE。数据先入 swap（工作态），回写窗口把桶写纪元与已同步纪元不等的脏桶序列化进 master（tmp+atomic move）；打开时删除 swap 并惰性加载 master。崩溃恢复点＝最后一次 master 同步：swap 因 DELETE_ON_CLOSE 不残留，master 永远原子且分段可验。", 10.5, False, "justify", 2, 6, 0)
table([
    ["主文件字节", "版本", "处理"],
    ["superblock=-0x200812250269L", "0x02", "旧整文件 ZSTD → tryParseBlinearV2 迁移入 swap → 下次 sync 升 v4"],
    ["同上", "0x03", "旧桶式 → 按桶懒加载迁移 → 升 v4（实测：真实上游 v0x03 文件全读回 + 原位升 v4）"],
    ["同上", "0x04", "本格式（读校验链：表哈希→桶哈希→段长→chunkSection 内部 xxhash32）"],
    ["superblock=0xc3ff13183cca9d9aL", "1/2/3", "xymb 祖宗线性 → 逐块迁移（保留）"],
    ["其他", "—", "硬错误拒绝（带超块十六进制与路径）"],
])
para("2.4　校验链与防挥写", 12, True, "left", 0, 6, 6)
para("加载：超块 → 版本分派 → 位置表（偏移边界 + footer xxhash64）→ 桶（len 上界 + 桶块 xxhash32，不解压即验）→ 段（长度上限）→ chunkSection（内部 xxhash32）。任一层失败 → IOException 携带文件与偏移。**任何加载失败都会置 loadFailed 标志，syncToMasterFile 拒绝执行**——防止 try-with-resources 关闭路径把空工作态写坏好的 master（实测捕获过该级联）。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 三 ----------------
para("三、集成接线", 14, True, "left", 0, 6, 12)
para("3.1　azurepatches-new（新类）", 12, True, "left", 0, 6, 6)
para("com.azurebranches.storage：BufferedLinearRegionFile（引擎，v4 改造）、BufferedLinearRegionFileFlusher（每文件保存调度/背压：nIo 线程 + 检查器 + 写后超时）、IRegionFile、RegionFormat、RegionFormatBootstrap；net.minecraft.world.level.chunk.storage：RegionFileAdapter（vanilla MCA 适配）。", 10.5, False, "justify", 2, 6, 0)
para("3.2　azurepatches-src 覆盖层", 12, True, "left", 0, 6, 6)
para("RegionFileStorage 整体覆盖：regionCache 与全部类型化局部与返回改为 IRegionFile；`new RegionFile(...)` 两处（getRegionFile / getRegionFileIfExists）改为 `RegionFormat.open(...)`；moonrise$ 接口方法签名同步。", 10.5, False, "justify", 2, 6, 0)
para("3.3　transformSource 锚点（build.gradle.kts）", 12, True, "left", 0, 6, 6)
para("ChunkSystemRegionFileStorage（moonrise$getRegionFileIfLoaded 等 2 处签名）、MoonriseRegionFileIO（regionFile 变量类型 + IORunnable.run(final RegionFile) 2 处）、ChunkSystemChunkBuffer（moonrise$write 参数）、RegionFile（ChunkBuffer.moonrise$write 参数）——全部 fail-fast 唯一匹配，任何上游漂移都会打断构建。依赖注入 folia 侧 build.gradle.kts：zstd-jni 1.5.4-1、zero-allocation-hashing 0.16；lz4-java 经 velocity-natives 已在运行时类路径。产物的 META-INF/libraries.list 已验证包含三项依赖。", 10.5, False, "justify", 2, 6, 0)
para("3.4　配置面", 12, True, "left", 0, 6, 6)
para("AzureBranchesConfig 新增 `storage.region_format`（「mca」 | 「b_linear_v4」，默认「mca」）与 `storage.linear.compression_level`（1..22，默认 1）；RegionFormatBootstrap 在首次区域文件创建前读取并应用（RegionFormat.open 内幂等调用），常见模块（azurebranches-common 只对 Folia API 编译）不引用存储引擎，选择逻辑落在 Folia 侧。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 四 ----------------
para("四、格式 v4 工程修复（F 系列）", 14, True, "left", 0, 6, 12)
table([
    ["修复", "问题", "做法"],
    ["F1 位置表翻转", "表/桶缓冲填满后 position==limit，writeFullyAt 写 0 字节（定位表与桶数据双处）", "写入前 flip()；桶哈希改在内存字节上计算"],
    ["F2 桶哈希跨度", "写侧哈希含 8B 长度头、读侧只哈希 zstd 载荷", "统一为整个桶块（rawLen|compressedLen|ZSTD），两侧一致（规范同步更新）"],
    ["F3 footer 读回", "write-only outChannel 上读回 table/bucket 字节导致 NonReadableChannelException", "改为内存字节数组计算 tableXXHash64 与桶哈希"],
    ["F4 v4 打开分派", "构造期 tryParseBlinearV2 视 v4 为非法版本", "v4 与 v3 同走惰性加载分派"],
    ["F5 防空写", "加载失败后 try-with-resources close 同步空工作态覆写 master", "loadFailed 标志 + syncToMasterFile 拒绝"],
])

# ---------------- 五 ----------------
para("五、构建与验证", 14, True, "left", 0, 6, 12)
para("5.1　构建管线", 12, True, "left", 0, 6, 6)
para("删除 folia-server/build/cache/folia-paperclip.jar 强制全量重建（buildFolia 的缓存跳过是已知陷阱：只动 azurepatches-new 时必须删缓存，否则旧类静默进产物）；`.\gradlew.bat :azurebranches-server:buildFolia :azurebranches-server:mergeJar --no-configuration-cache`，产物 azurebranches-server-26.1.2-AB-0002-EXP6Plus.jar（57 MB）。构建日志确认全部锚点命中、编译零错误。", 10.5, False, "justify", 2, 6, 0)
para("5.2　JVM 级 harness", 12, True, "left", 0, 6, 6)
para("环境：folia-src 平铺服务器 jar + 运行时 libraries（107 项）+ 编译期注解；各阶段先 SharedConstants.tryDetectVersion() + Bootstrap.bootStrap()（绕过“Game version not set”）；20 次构建迭代中排除的测试环境问题：paperclip jar 无普通 class 条目（改用平铺 jar）、cmd 8191 字符限制（扁平 cpdir + 通配符）、logging-1.6.11.jar 遮蔽服务端 LogUtils（剔除）、NBT 26.1 Optional 返回值与 CompoundTag final 化（harness 适配）。", 10.5, False, "justify", 2, 6, 0)
table([
    ["用例", "内容", "结果"],
    ["round-trip", "47 chunk（桶 0 全 32 + 每桶 1）写→关→开→读全验；覆盖写 (3,7) 保留；邻桶完好", "PASS（186ms/47 chunk；master 3.08MB；表/桶哈希互验通过）"],
    ["损坏 2a", "桶数据区翻字节", "PASS（读侧拒绝，不返回部分数据）"],
    ["损坏 2b", "位置表翻字节（footer xxhash64 层）", "PASS"],
    ["损坏 2c", "截断 footer", "PASS"],
    ["kill-test", "先写 1 已提交 chunk；子 JVM 写 48 chunk 后硬杀；重开", "PASS（已提交 chunk 完好；swap 不残留）"],
    ["v3 读兼容+迁移", "真实上游引擎（Arbor-ver-26.1.2 源码原样编译）生成 v0x03；我方引擎读 47 全对 → 追加写 (5,5) → 原位升 v4 → 重读 47+1 全对", "PASS"],
    ["性能基线", "64×64KiB 写+flush：v4 44.8 MiB/s vs MCA 12.5 MiB/s；v4 master 4.20MB vs MCA 4.46MB", "PASS（≈3.6×，体积略优）"],
])
para("5.3　服务端级冒烟（全新存档）", 12, True, "left", 0, 6, 6)
para("隔离目录 exp7-test（server-port 25570 / rcon 25576 / level-name=world-exp7-v4）；配置文件写入 [storage] region_format=「b_linear_v4」。日志断言：`[AzureBranches] storage.region_format=b_linear_v4 (compression 1)` 且全程无 Exception/ERROR；启动 Done(29.8s) → 生成区写入；region 目录出现 r.*.mca（超块+version=0x04 逐字节验证）与 .swp；优雅 stop 全 I/O 完成；重启 Done(20.9s) 无错；forceload 若干区后硬 kill（Stop-Process -Force）→ master 尺寸/头不变、无 swap 残留；再重启 Done(18.3s) 无错；**三维度（overworld/nether/the_end）+ entities 共 14 个主文件全部 version=4**（entities 文件同走 RegionFormat.open，确认抽象覆盖实体数据）。", 10.5, False, "justify", 2, 6, 0)

# ---------------- 六 ----------------
para("六、已知限制与后续工作", 14, True, "left", 0, 6, 12)
para("6.1　已知限制", 12, True, "left", 0, 6, 6)
para("1. **崩溃恢复点＝最后一次 master 同步**：swap 为 DELETE_ON_CLOSE 工作文件，硬杀后 swap 不残留（上游语义，忠于移植）；写后超时（默认 3s）内的未同步写会丢失——一致性以“master 永不被部分写/空写污染”保证。", 10.5, False, "justify", 2, 3, 0)
para("2. **MCA ↔ b_linear 不互通**：切到 b_linear_v4 后旧 MCA 区域文件无法解码（v4 读兼容只覆盖 xymb 系）；老世界如需启用需先按 xymb 迁移路径或整体重置，配置变更应配合世界备份。", 10.5, False, "justify", 2, 3, 0)
para("3. v0x02（整文件 ZSTD）读路径由上源 tryParseBlinearV2 保留，本次以真实上游样本验证的是 v0x03（上游当前写版本）；v0x02 仅由代码路径保证，未生成样本实测。", 10.5, False, "justify", 2, 3, 0)
para("4. 性能为单机微基准（同进程顺序写），未覆盖真实负载下的多区域并发、长时段混合读写与磁盘压力曲线。", 10.5, False, "justify", 2, 3, 0)
para("6.2　后续工作（优先级排序）", 12, True, "left", 0, 6, 6)
para("**P0**：v0x02 样本实测；MCA→v4 单区域迁移工具（一次性转换 + 校验）；服务端压力基线（并发 region / 大世界跑图）；写后 3s 超时参数验证（短超时=高损窗口）。", 10.5, False, "justify", 2, 3, 0)
para("**P1**：sync 失败重试/告警出口（当前 flusher 日志 + loadFailed 拒绝已保证不静默，补 metrics 计数）；压缩级别按世界特征自动选择。", 10.5, False, "justify", 2, 3, 0)
para("**P2**：把 swap 改为可选持久 WAL（保留 crash 前未同步写的恢复能力）——与上游语义偏离需单独评估兼容性。", 10.5, False, "justify", 2, 3, 0)

# ---------------- 七 ----------------
para("七、结论", 14, True, "left", 0, 6, 12)
para("EXP7 让 AzureBranches 具备可切换、可校验、可迁移的区域文件存储后端：默认 MCA 完全无行为漂移，切到 b_linear_v4 即获得读写放大优化与四层校验（位置表 xxhash64 / 桶块 xxhash32 / 段长上限 / chunkSection 内部哈希）与“加载失败拒覆写”的防损坏语义；v0x03 真实上游样本的全读回与原位升 v4 验证了“老世界就地升级”路径；JVM harness 与服务端冒烟（全新存档、三维度 + entities、优雅停/重启、硬杀恢复）共 7 组用例全绿。", 10.5, False, "justify", 2, 6, 0)
para("方法学上，本版本再次验证：测试环境事实必须显式建模（paperclip 无 class 条目、cmd 限长、日志遮蔽、NBT 26.1 API、旧类缓存陷阱）；分层跟踪（文件字节级 HashCheck/BucketProbe 与引擎 diag 相结合）把一个“迁移后新 chunk 缺失”的伪快速定位为上位机 verify 的 blob 配方不一致，而非引擎缺陷——字节级证据先于推论。", 10.5, False, "justify", 2, 6, 0)
para("文档：AzureBranches-26.1.2-EXP7　　版本：26.1.2-EXP7　　日期：2026 年 08 月 22 日　　项目：AzureBranches (https://github.com/XCxyTianQ/AzureBranches)", 9, False, "center", 0, 12, 0)

doc.save(OUT)
print("DOCX written:", OUT)
